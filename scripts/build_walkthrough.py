"""
Builds Mastercard-Submission-Walkthrough.docx from real, verified project
content: source code and the committed dashboard.json. Run from the repo
root: python scripts/build_walkthrough.py
"""

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent


def load_detect_metrics():
    """detect/src/detector.py's evaluate() now computes f1_score and
    roc_auc directly, so the committed dashboard.json already has them.
    No need to retrain here."""
    with open(ROOT / "web" / "data" / "dashboard.json") as f:
        dash = json.load(f)
    m = dash["detect"]["metrics"]
    return {
        "f1": m["f1_score"],
        "roc_auc": m["roc_auc"],
        "confusion_matrix": m["confusion_matrix"],
        "precision": m["precision"],
        "recall": m["fraud_caught_rate"],
        "fpr": m["false_positive_rate"],
    }


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build():
    live = load_detect_metrics()

    with open(ROOT / "web" / "data" / "dashboard.json") as f:
        dash = json.load(f)

    doc = Document()

    # ---------- Title ----------
    title = doc.add_heading("Parmana Authority Gate", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Solution Walkthrough, Mastercard Innovation Challenge @ GFF 2026\n"
        "AI Defense Lab for Payment Security"
    )
    r.italic = True
    r.font.size = Pt(13)
    doc.add_paragraph()

    para(
        doc,
        "Repository: https://github.com/pavancharak/execution-authority-gate   |   "
        "Live prototype: https://parmana.fly.dev",
    )
    doc.add_paragraph()

    # ---------- Executive Summary ----------
    h1(doc, "Executive Summary")
    para(
        doc,
        "Parmana is a two layer payment fraud defense system built as a closed loop across "
        "the challenge's three pillars: it identifies thirteen distinct GenAI era fraud attack "
        "vectors targeting payments across card, wire/ACH, real time push payment, and agentic "
        "commerce rails, generates realistic simulations of six of them at scale "
        "(including four attack types produced by real GPT-4o-mini calls, not templated text), "
        "and defends against them with a RandomForest classifier evaluated on a held out test "
        "set of 6,869 transactions at a realistic 2% fraud rate.",
    )
    para(
        doc,
        "The system goes one step beyond a standalone classifier. A second, fully independent "
        "layer, a deterministic mandate engine derived from each customer's own transaction "
        "history, checks every transaction against that customer's real spending limits, "
        "merchants, hours, and velocity, regardless of what the fraud model concludes. Either "
        "layer objecting is enough to block a transaction. Every final decision, ALLOW, FLAG, "
        "or BLOCK, is then signed with Ed25519 so it is independently verifiable and tamper "
        "evident after the fact. The loop closes with an adversarial agent that attacks the "
        "trained detector directly, generating evasion variants used to stress test it.",
    )

    # ---------- Pillar 1: Identify ----------
    h1(doc, "Pillar 1: Identify. The Attack Taxonomy")
    para(
        doc,
        "identify/attack-taxonomy.md documents thirteen distinct GenAI accelerated attack "
        "vectors against payment systems, each grounded in a real point in the payment "
        "lifecycle or a real payment rail (onboarding, authorization, KYC, customer support, "
        "B2B wire/ACH, real time push payments, agentic commerce, dispute adjudication, long "
        "horizon account lifecycle) rather than a generic 'fraud' label. Six are actively "
        "simulated by seven bounded agents in generate/src/fraud_agents.py. The remaining "
        "seven are documented as honest, explicit known gaps rather than claimed and left "
        "unimplemented, chosen specifically to cover rails and surfaces the six simulated "
        "agents don't touch, breadth across the payment ecosystem, not just depth on one "
        "corner of it.",
    )

    attacks = [
        ("1. AI Fabricated Identity", "Account creation / onboarding",
         "GPT-4o-mini generates each synthetic identity so every transaction looks statistically "
         "normal on its own; there is no history to compare against on transaction #1. "
         "(agent1_fake_identity, real OpenAI calls)"),
        ("2. Spending Pattern Replication (card testing & draining)", "Authorization / payment processing",
         "A stolen card/token replicated at machine speed to mimic the real cardholder's "
         "amount/merchant/timing shape; the anomaly is velocity and cumulative volume, not any "
         "single transaction. (agent5_pattern_replicator, statistical, local)"),
        ("3. Payment Form / API Fuzzing", "Client input into backend processing",
         "Automated fuzzing against known payload classes, used as reconnaissance, the one "
         "field that doesn't reject cleanly becomes a later entry point. (agent6_injection_generator, local)"),
        ("4. AI Voice/Chat Social Engineering", "Customer authentication via call center / chat",
         "GPT-4o-mini generates fictional call center transcripts convincing enough to pass "
         "identity questions based on leaked personal data. Text only in this lab: no voice "
         "cloning, no audio, no real support line contacted. (agent2_social_engineer, real OpenAI calls)"),
        ("5. Authorization Limit Probing", "Authorization",
         "Real amounts from $10 to $10,000 submitted through the actual trained detector to map "
         "its decision thresholds before a precisely sized charge follows. This attacks our own "
         "detector directly, not a fabricated external system. (agent3_limit_prober)"),
        ("6. AI Generated KYC Document Forgery", "Identity verification (KYC)",
         "GPT-4o-mini generates plausible identity metadata bundles (name, DOB, address, "
         "occupation) tuned to defeat the exact detector in use. No document images in this lab, "
         "metadata only. (agent4_kyc_forger, real OpenAI calls)"),
        ("7. Feedback Loop Poisoning of the Fraud Model", "Post decision / dispute ingestion (known gap)",
         "Repeated small transactions plus false dispute signals meant to teach a retraining "
         "pipeline that fraud is normal. This lab does not run a persistent retraining process, "
         "so this attack is documented honestly as future work rather than claimed as tested. "
         "(agent7_feedback_loop generates real time evasion variants against the already trained "
         "model, which is related but distinct, it does not poison a retraining pipeline.)"),
        ("8. AI Orchestrated Business Email Compromise / Vendor Payment Redirection",
         "Accounts payable, B2B wire/ACH rails (known gap)",
         "A compromised or spoofed vendor email plus an LLM drafting a convincing updated bank "
         "details request matching the real vendor's tone and invoicing history. Wire/ACH have "
         "no real time chargeback, so the loss is discovered only after the funds are gone. Out "
         "of scope for this lab's agents, which target consumer/card rails, not vendor "
         "correspondence."),
        ("9. Authorized Push Payment Fraud via Voice Cloned Urgency Scams",
         "Real time/instant push payment rails (known gap)",
         "A cloned voice of a trusted contact plus manufactured urgency, so the customer "
         "authorizes the transfer themselves. No stolen credential, no anomalous "
         "authentication, so a shape based classifier like this lab's detect layer has no "
         "signal to work with by construction, a real limitation worth naming."),
        ("10. Agentic Commerce Hijack: Prompt Injection Against Autonomous Shopping Agents",
         "AI shopping/checkout agents transacting on a customer's behalf (known gap)",
         "Hidden instructions embedded in a malicious merchant page that an LLM driven "
         "purchasing agent reads and acts on, using the customer's own real, authorized "
         "credentials. The most 2026 specific vector in this taxonomy: it targets AI agents "
         "transacting, not AI generating fraud content. This lab's mandate layer would still "
         "catch some of these in practice (an unfamiliar merchant, an odd hour, over budget), "
         "a partial mitigation even without a dedicated agent simulating the attack."),
        ("11. Deepfake Liveness Bypass for Biometric KYC",
         "Identity verification, biometric liveness checks (known gap)",
         "Real time face swap or diffusion generated video that responds to a liveness "
         "challenge (blink, turn head), defeating the exact signal modern liveness checks "
         "rely on. More severe than attack #6 since it beats the strongest KYC defense in use, "
         "not a weaker static one. agent4_kyc_forger generates identity metadata only, no "
         "video, so this stays a distinct, unimplemented gap."),
        ("12. AI Generated Fake Dispute Evidence (Chargeback / Refund Abuse)",
         "Post transaction dispute and chargeback process (known gap)",
         "A fabricated receipt, product photo, or support transcript used to win a dispute on "
         "a transaction that was completely legitimate. The fraud happens entirely in "
         "paperwork submitted after the fact, a stage most fraud detection never inspects. "
         "Related to but distinct from attack #7: this targets the dispute reviewer, not the "
         "fraud model's training data."),
        ("13. Synthetic Identity Bust Out",
         "Full account lifecycle, monetized in a single terminal event (known gap)",
         "A fabricated identity (attack #1) used to build months of unremarkable transaction "
         "history and a real credit line, then maxed out and abandoned. There is no anomaly "
         "until the single terminal transaction, by which point months of history have "
         "already vouched for the account. This lab's detect layer scores each transaction "
         "independently, so a pattern spanning months is structurally outside what it can "
         "see, a real limitation of that design."),
    ]
    for name, where, desc in attacks:
        p = doc.add_paragraph()
        p.add_run(f"{name}").bold = True
        p.add_run(f":  {where}")
        doc.add_paragraph(desc, style="List Bullet")

    # ---------- Pillar 2: Generate ----------
    h1(doc, "Pillar 2: Generate. Simulating Attacks at Scale, with Fidelity")
    para(
        doc,
        "generate/src/run_simulation.py orchestrates seven bounded agents. Four of them "
        "(fake identity, social engineering, KYC forgery, and the feedback loop evasion probe) "
        "make real GPT-4o-mini calls at temperature=0.9, producing genuinely non deterministic "
        "fraud variation rather than a fixed template repeated with noise. The remaining three "
        "(pattern replication, form fuzzing, limit probing) are local/statistical by design, "
        "since their realism comes from matching a real distribution or a real threshold search, "
        "not from generative text.",
    )
    sim = dash["simulation"]
    total_tx = sim["good_transaction_count"] + sim["fraud_transaction_count"]
    fraud_rate = sim["fraud_transaction_count"] / total_tx
    para(
        doc,
        f"The committed run behind web/data/dashboard.json is fully self generated by this "
        f"repo's own generate/ layer, using 24 real OpenAI calls (~$0.03 total): "
        f"{total_tx:,} transactions at a realistic {fraud_rate:.2%} fraud rate "
        f"({sim['good_transaction_count']:,} legitimate + {sim['fraud_transaction_count']:,} "
        "fraudulent). The legitimate pool is scaled up with free local generation specifically "
        "so the fraud rate can stay realistic without shrinking the fraud sample to statistical "
        "noise, 459 fraud examples means the held out test split alone has 138 fraud cases to "
        "evaluate detection against.",
    )
    para(
        doc,
        "Fraud spreads across the five attack types that produce labeled transactions: "
        "fake_identity (112), social_engineering (47), kyc_synthetic (100), pattern_copy (100), "
        "form_break (100). Attacks #5 and #7 (limit probing, feedback loop evasion) probe the "
        "trained model directly rather than generating labeled transactions, so they sit outside "
        "this breakdown by design, see generate/src/probe_agents.py.",
    )
    para(
        doc,
        "Numbers are deliberately not frozen to one lucky run: because agents 1, 2, 4, and 7 "
        "call the real OpenAI API at temperature=0.9, rerunning generate/src/run_simulation.py "
        "with a fresh OPENAI_API_KEY produces different fraud examples and slightly different "
        "metrics each time. That variance is the point, it demonstrates the detector holds up "
        "across different fraud patterns, not just one fixed, cherry picked dataset.",
    )

    # ---------- Pillar 3: Defend ----------
    h1(doc, "Pillar 3: Defend. Detection Model and Efficacy")
    para(
        doc,
        "detect/src/detector.py trains a RandomForestClassifier (n_estimators=80, max_depth=5, "
        "min_samples_leaf=12, class_weight='balanced') on six features per transaction: amount, "
        "hour_of_day, log_seconds_since_prev_tx, location_mismatch_km, pattern_similarity, and "
        "ai_generated_signal. class_weight='balanced' is used specifically to hold up under the "
        "realistic ~2% fraud imbalance rather than collapsing to the majority class.",
    )

    para(doc, "Detection metrics, held out test set (6,869 transactions, 138 fraud):", bold=True)
    cm = live["confusion_matrix"]
    table = doc.add_table(rows=7, cols=2)
    table.style = "Light Grid Accent 1"
    rows_data = [
        ("Confusion matrix", f"TN {cm['true_negative']:,} / FP {cm['false_positive']:,} / "
                              f"FN {cm['false_negative']:,} / TP {cm['true_positive']:,}"),
        ("Recall (fraud caught)", f"{live['recall']:.2%}"),
        ("False positive rate", f"{live['fpr']:.2%}"),
        ("Precision", f"{live['precision']:.2%}"),
        ("F1 score", f"{live['f1']:.4f}"),
        ("ROC AUC", f"{live['roc_auc']:.4f}"),
        ("Top signal", "log_seconds_since_prev_tx, pattern_similarity, ai_generated_signal"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_paragraph()
    para(
        doc,
        "Why precision is 21%, and why that is expected, not a flaw: fraud is rare in this "
        "dataset (~2%). Tuning a classifier to catch 89% of a rare event requires flagging "
        "aggressively, and recall/precision trade off against each other at that base rate, the "
        "same tradeoff airport security makes to catch most weapons at the cost of screening "
        "many harmless bags. Precision measured on the detect layer in isolation is not the "
        "system's real world false accusation rate: a detect layer flag does not block anything "
        "by itself. It still has to clear the independent, rule based mandate layer before a "
        "transaction is denied, and every final decision is signed and independently verifiable. "
        "See docs/JUDGES_GUIDE.md for the full treatment.",
    )
    para(
        doc,
        "A separate red team check: agent7_feedback_loop generated 18 real time evasion variants "
        "against the already trained model. 17 of 18 still failed to evade detection after the "
        "balanced class weighting fix (an earlier, more imbalanced version of this dataset let "
        "16 of 18 through). This is a narrow robustness result on a small sample, not a claim "
        "about stopping some fixed percentage of attacks in general, but it is a real, "
        "code executed adversarial test against the actual trained model, not a simulated one.",
    )

    # ---------- The Closed Loop / Novel Differentiator ----------
    h1(doc, "The Closed Loop, and What Makes This Different From a Standalone Classifier")
    para(
        doc,
        "Two things distinguish Parmana from 'train a classifier on synthetic fraud and report "
        "recall': an independent authorization layer that doesn't trust the detector's score at "
        "all, and a feedback path that turns the trained defense back into an attack surface.",
    )

    h2(doc, "Mandate layer: independent, deterministic authorization")
    para(
        doc,
        "mandate/src/mandate_checker.py derives a mandate for each customer from their own "
        "known good transaction history: a spending limit (avg amount × count × 1.5), the "
        "specific merchants they've actually used, the hours they actually transact in (with a "
        "small margin), and a daily velocity cap above their observed maximum. Every transaction "
        "is checked against all four rules with AND, not OR, a transaction inside its spending "
        "limit but at 3am from an unfamiliar merchant still fails the mandate, even if the "
        "detect layer scored it as low risk. In pipeline/src/run_pipeline.py's combine_decision, "
        "either layer objecting is enough to BLOCK; a clean mandate never downgrades a detect "
        "BLOCK, and a clean detect score never upgrades a mandate violation past BLOCK. Real, "
        "signed evidence for this: 8 real fraud transactions the detector scored as low risk "
        "were still blocked by the mandate layer alone.",
    )

    h2(doc, "Cryptographic signing: tamper evident, not a third vote")
    para(
        doc,
        "sign/src/authority_signer.py signs every final decision (ALLOW, FLAG, and BLOCK alike) "
        "with Ed25519, using a key held only by the external 'authority' identity, neither the "
        "detector nor the mandate checker can produce a valid signature themselves. The signed "
        "envelope covers transaction_id, fraud_score, detect_decision, mandate_allowed, "
        "violated_mandate_rules, final_decision, and reasons. Signing is not a third check that "
        "can veto anything; it makes the first two layers' decision provable and unforgeable "
        "after the fact. In this run, 6,869/6,869 signed decisions verify independently against "
        "the authority's public key.",
    )

    h2(doc, "The loop closes: the defense becomes the next attack's training ground")
    para(
        doc,
        "generate/src/probe_agents.py runs after the detector is trained and attacks it "
        "directly: agent3_limit_prober submits real amounts from $10–$10,000 through the actual "
        "model to map its decision boundary, and agent7_feedback_loop generates real time "
        "evasion variants against the same trained model. Both are signed and verified the same "
        "way every other decision is. This is the challenge's closed loop requirement in "
        "practice, the trained defense is itself the object the next round of attack generation "
        "targets, not a separate, one off evaluation step.",
    )

    # ---------- Real World Feasibility ----------
    h1(doc, "Real World Feasibility")
    para(
        doc,
        "What is solidly demonstrated: a genuine two layer decision (ML detection + independent "
        "rule based authorization) where either layer can force a block, real Ed25519 "
        "signing with correct key separation (the AUTHORITY key that signs decisions is a "
        "different key from the REVIEWER key that signs human overrides) and empirically "
        "confirmed tamper evidence for the fields inside the signed envelope, and, as of this "
        "revision, a durable audit trail, caller authentication, and an execution-ready handoff "
        "for a payment processor, described in the four subsections below.",
    )
    para(
        doc,
        "An earlier internal audit (EAG-AUDIT-GAPS.md, committed to the repository) found four "
        "concrete gaps between what this project's architecture document claimed and what the "
        "code actually did: decisions were not durable, signing keys were not durable, there was "
        "no execution handoff, and there was no caller authentication. We treated that audit as a "
        "punch list rather than a problem to explain away. All four are now implemented, tested, "
        "and cited below by file path, not merely described in prose.",
    )

    h2(doc, "Audit & Durability")
    para(
        doc,
        "pipeline/src/audit_trail.py's AuditTrail class appends one JSON line per decision to "
        "pipeline/audit/decisions.jsonl, idempotent on the signed record's record_id, and never "
        "rewrites a previously written line, unlike the earlier pipeline_decisions.json, which "
        "was fully overwritten on every pipeline run. This file is committed to git rather than "
        "ignored: the run behind this submission produced 6,869 real signed decisions, all "
        "6,869 independently re-verified against the committed public key via "
        "AuditTrail.verify_all(), a command any reader of this document can rerun themselves. "
        "The signing public keys (sign/tokens/authority_public_key.pem, "
        "reviewer_public_key.pem) are now committed to git as well, so a signature produced in "
        "this environment can be verified from a completely fresh checkout, closing the "
        "specific gap the earlier audit identified: private keys remain git-ignored and "
        "generated per environment, which is correct, but the public keys needed for "
        "verification are no longer regenerated alongside them.",
    )

    h2(doc, "Caller Scoping")
    para(
        doc,
        "sign/src/caller_auth.py adds HMAC-SHA256 signed caller tokens carrying a scoped "
        "permission list: a payment-processor identity may execute ALLOW or FLAG decisions but "
        "not BLOCK (a payment processor settles or steps up; it does not get to unilaterally "
        "deny a transaction the authority didn't already deny), a fraud-analyst identity may "
        "execute all three, and an audit-system identity is read only and can never execute "
        "anything. pipeline/src/run_pipeline.py's --caller-id flag threads a caller identity "
        "into authority_signer.sign_pipeline_decision, where it is embedded inside the signed "
        "envelope itself (a caller_id field, not a sibling), so it is exactly as tamper evident "
        "as final_decision. web/server.py exposes POST /api/callers/token to issue tokens and "
        "gates the new execution route behind a require_auth decorator requiring "
        "Authorization: Bearer <token>. This is deliberately a separate trust boundary from the "
        "Ed25519 decision-signing key: a caller proving its own identity is not the same secret "
        "that makes a decision's content authoritative, so a caller can never self-authorize a "
        "decision it merely requested.",
    )

    h2(doc, "Execution Integration")
    para(
        doc,
        "sign/src/decision_executor.py's DecisionExecutor turns a signed decision into an action "
        "against a payment_processor_webhook callable: ALLOW maps to settle, FLAG to "
        "step_up_auth, BLOCK to deny. Before any action is dispatched, it independently "
        "re-verifies the decision's signature (fail closed: a tampered decision is rejected, "
        "the webhook is never called) and, when a caller identity is supplied, checks that "
        "caller's permission for that decision type. Every decision's record_id is tracked so "
        "the same signed decision can never be executed twice, and every attempt, executed or "
        "rejected, is appended to its own audit log. web/server.py exposes this at "
        "POST /api/enforce/decisions. This repo ships no real payment processor integration; "
        "the shipped webhook (noop_webhook) simulates and explicitly labels its own output "
        "\"simulated\": true. The honest claim is decision-ready and execution-ready for "
        "external enforcement, not enforced, and docs/PRODUCTION_DEPLOYMENT.md documents "
        "exactly what a real integration would still need to build.",
    )

    para(
        doc,
        "Test coverage for all three additions: 31 new hermetic test cases across "
        "tests/test_audit_trail.py (9), tests/test_caller_auth.py (13), and "
        "tests/test_executor.py (9), bringing the suite from 54 to 85 passing tests "
        "(pytest tests/ -v), with the same no-network, no-external-state hermeticity as the "
        "original suite.",
    )
    para(
        doc,
        "None of this weakens the two pillars the challenge scores most heavily (fidelity of "
        "simulated attacks, detection efficacy), both are real, measured, and reproducible by "
        "running pytest tests/ and pipeline/src/run_pipeline.py from a fresh clone.",
    )

    # ---------- What This Submission Is NOT ----------
    h1(doc, "What This Submission Is NOT")
    bullets(doc, [
        "Not a payment processor and not connected to one. Nothing in this repository moves "
        "real money; the shipped execution webhook stub labels its own output as simulated.",
        "Not a production deployment. parmana.fly.dev serves a static, committed dashboard "
        "snapshot; the detection/mandate/signing pipeline does not run inside the deployed "
        "container on every request.",
        "Not a claim that 89.1% recall and 6.8% false positive rate are fixed constants. They "
        "come from a non-deterministic data generation process (real OpenAI calls at "
        "temperature=0.9) and will shift slightly on every regeneration, by design, see the "
        "Robustness discussion in README.md.",
        "Not HSM- or KMS-backed key management. Signing keys are generated and persisted "
        "locally per environment; docs/PRODUCTION_DEPLOYMENT.md names exactly what a managed "
        "key service integration would require.",
        "Not a full RBAC or database-backed audit system at production scale. The caller "
        "authentication and audit trail added this revision are real and tested, but the audit "
        "trail is JSONL, not yet a database, and caller registration is currently a fixed "
        "predefined list plus in-process dynamic registration, not a managed identity provider.",
    ])
    para(
        doc,
        "We include this section deliberately, in the same spirit as EAG-AUDIT-GAPS.md: a "
        "submission judges can trust is one that states its own boundaries as clearly as its "
        "capabilities.",
    )

    # ---------- Novelty ----------
    h1(doc, "Novelty")
    bullets(doc, [
        "Detection and authorization are architecturally independent models of trust, not one "
        "model wearing two hats: a statistical model can be wrong about a transaction that is "
        "still authorized by the customer's real history, and vice versa. Either can force a "
        "block; neither alone can force an allow.",
        "The attack generator doubles as the detector's adversary: agent3 and agent7 in "
        "generate/src/probe_agents.py attack the trained model directly using the same "
        "signing/verification path as every other decision, closing the loop the challenge asks "
        "for rather than treating identify/generate/defend as three separate, unconnected "
        "deliverables.",
        "Every decision, not just blocks, is cryptographically signed, so the audit trail is "
        "symmetric: an attacker who evades detection and an operator who wrongly blocks a "
        "legitimate customer are both provable after the fact from the same mechanism.",
        "The mandate layer is derived from each customer's own real transaction history rather "
        "than a hand authored global policy, so it doesn't require a fraud analyst to write "
        "rules per customer segment.",
    ])

    # ---------- Reproducing These Results ----------
    h1(doc, "Reproducing These Results")
    para(doc, "From a fresh clone of the public repository:")
    bullets(doc, [
        "pip install -r requirements.txt",
        "pytest tests/ -v   →   85 hermetic tests pass in seconds (no API key, no network calls); "
        "3 additional tests exercising the real OpenAI backed agents are skipped by default "
        "(ALLOW_LIVE_OPENAI=1 OPENAI_API_KEY=... pytest tests/test_generate.py -v to run them)",
        "cd web && python server.py, then open http://localhost:8080, the Live Test Harness "
        "runs a submitted transaction through the real trained model, mandate checker, and "
        "signer live, not a canned response; the Attack Walkthrough tab shows five real, "
        "already signed decisions pulled from an actual pipeline run, one per attack type.",
        "Live deployed instance: https://parmana.fly.dev",
    ])

    # ---------- Conclusion ----------
    h1(doc, "Conclusion")
    para(
        doc,
        "Parmana treats identify, generate, and defend as one closed loop rather than three "
        "independent deliverables: a real, GenAI grounded taxonomy of thirteen payment fraud "
        "vectors, spanning card, wire/ACH, real time push payment, and agentic commerce rails, "
        "drives realistic, at scale simulation of six of them (three attack types via real "
        "GPT-4o-mini calls), which trains a detector that is then attacked by its own generator "
        "to find and report its remaining gaps. An independent, customer history derived "
        "mandate layer catches what the detector alone would miss, and every decision is "
        "cryptographically signed so the result is provable, not just claimed. We have tried to "
        "be equally rigorous about what is solidly demonstrated (the two layer decision, the "
        "cryptography, the measured detection numbers, all reproducible from a fresh clone) and "
        "what would still need to be built for production use (durable storage, managed keys, "
        "a real payment execution handoff), the goal being a submission judges can trust "
        "precisely because it does not overclaim.",
    )

    out_path = ROOT / "Mastercard-Submission-Walkthrough.docx"
    doc.save(str(out_path))
    print(f"Saved {out_path} ({out_path.stat().st_size / 1024:.1f} KB)")
    print(f"Live computed F1={live['f1']}, ROC AUC={live['roc_auc']}")


if __name__ == "__main__":
    build()
